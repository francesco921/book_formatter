# book_formatter.py
# UI Streamlit + CLI. Genera DOCX e PDF (nativo, senza LibreOffice) con stile editoriale.
# Formati pagina: 6x9 e 8.5x11. TOC per sezioni (SECTION I/II/...) e capitoli (CHAPTER N — ...).

import argparse
import datetime
import os
import re
import sys
from typing import List, Tuple, Optional

# ----------------------------
# Tipi
# ----------------------------
HeadingItem = Tuple[int, str]        # (livello, testo) 1=Sezione, 2=Capitolo
ContentBlock = Tuple[int, List[str]] # (livello, [paragrafi])

# ----------------------------
# DOCX (costruzione output)
# ----------------------------
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------------------
# PDF (input parsing opzionale)
# ----------------------------
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
except Exception:
    pdf_extract_text = None

# ----------------------------
# PDF (output nativo)
# ----------------------------
from reportlab.lib.units import cm, inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, PageBreak, Spacer
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle


# =====================================================================
# PARSING INPUT (DOCX/PDF) -> headings + content_blocks
# Riconosce:
#   - Stili Heading 1/Heading 2 (o Titolo 1 / Titolo 2)
#   - Pattern numerati "SECTION ..." e "CHAPTER \d+" oppure "1." / "1.1 ..."
# =====================================================================
def parse_docx_input(path: str) -> Tuple[List[HeadingItem], List[ContentBlock]]:
    doc = Document(path)
    headings: List[HeadingItem] = []
    blocks: List[ContentBlock] = []
    curr_level: Optional[int] = None
    curr_buf: List[str] = []

    def flush():
        nonlocal curr_level, curr_buf
        if curr_level is not None:
            blocks.append((curr_level, curr_buf))
        curr_level = None
        curr_buf = []

    # Pattern alternativi
    sec_pat = re.compile(r"^\s*SECTION\s+([IVXLC]+)\b.*", re.I)
    ch_pat  = re.compile(r"^\s*CHAPTER\s+(\d+)\b.*", re.I)
    h1_pat  = re.compile(r"^\s*\d+\.\s+.+")
    h2_pat  = re.compile(r"^\s*\d+\.\d+\s+.+")
    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue
        style = (p.style.name if p.style is not None else "")
        level: Optional[int] = None

        if style in ("Heading 1", "Titolo 1"):
            level = 1
        elif style in ("Heading 2", "Titolo 2"):
            level = 2
        else:
            if sec_pat.match(text) or h1_pat.match(text):
                level = 1
            elif ch_pat.match(text) or h2_pat.match(text):
                level = 2

        if level:
            flush()
            headings.append((level, text))
            curr_level = level
        else:
            if curr_level is None:
                curr_level = 1
                headings.append((1, "SECTION I — INTRODUCTION"))
            curr_buf.append(text)

    flush()
    return headings, blocks


def parse_pdf_input(path: str) -> Tuple[List[HeadingItem], List[ContentBlock]]:
    if pdf_extract_text is None:
        raise RuntimeError("pdfminer.six non installato. Esegui: pip install pdfminer.six")
    raw = pdf_extract_text(path)
    lines = [ln.strip() for ln in raw.splitlines() if ln.strip()]

    headings: List[HeadingItem] = []
    blocks: List[ContentBlock] = []
    curr_level: Optional[int] = None
    curr_buf: List[str] = []

    def flush():
        nonlocal curr_level, curr_buf
        if curr_level is not None:
            blocks.append((curr_level, curr_buf))
        curr_level = None
        curr_buf = []

    sec_pat = re.compile(r"^\s*SECTION\s+([IVXLC]+)\b.*", re.I)
    ch_pat  = re.compile(r"^\s*CHAPTER\s+(\d+)\b.*", re.I)
    h1_pat  = re.compile(r"^\s*\d+\.\s+.+")
    h2_pat  = re.compile(r"^\s*\d+\.\d+\s+.+")
    for ln in lines:
        if sec_pat.match(ln) or h1_pat.match(ln):
            flush()
            headings.append((1, ln))
            curr_level = 1
        elif ch_pat.match(ln) or h2_pat.match(ln):
            flush()
            headings.append((2, ln))
            curr_level = 2
        else:
            if curr_level is None:
                curr_level = 1
                headings.append((1, "SECTION I — INTRODUCTION"))
            curr_buf.append(ln)

    flush()
    return headings, blocks


# =====================================================================
# "CLEAN" HEURISTIC: normalizza l’estrazione in stile desiderato
# - Sezione -> "SECTION I — TITOLO"
# - Capitolo -> "CHAPTER N — Titolo"
# =====================================================================
_ROMANS = ["I","II","III","IV","V","VI","VII","VIII","IX","X","XI","XII","XIII","XIV","XV","XVI","XVII","XVIII","XIX","XX"]

def _roman(n: int) -> str:
    if 1 <= n <= len(_ROMANS):
        return _ROMANS[n-1]
    # fallback semplice
    return "I"

def clean_structure(headings: List[HeadingItem], blocks: List[ContentBlock]) -> Tuple[List[HeadingItem], List[ContentBlock]]:
    cleaned_h: List[HeadingItem] = []
    sec_idx = 0
    ch_idx  = 0
    for level, text in headings:
        t = " ".join(text.replace("—","-").split()).strip(" .")
        if level == 1:
            sec_idx += 1
            # prendi parte dopo eventuale "SECTION *"
            m = re.search(r"section\s+[IVXLC]+\s*[—-]\s*(.+)", t, re.I)
            title = (m.group(1) if m else re.sub(r"^\d+\.\s*", "", t, flags=re.I))
            title = title.upper()
            t = f"SECTION {_roman(sec_idx)} — {title}"
        else:
            ch_idx += 1
            # prendi parte dopo eventuale "CHAPTER N"
            m = re.search(r"chapter\s+\d+\s*[—-]\s*(.+)", t, re.I)
            title = (m.group(1) if m else re.sub(r"^\d+\.\d+\s*", "", t, flags=re.I))
            # Capitalizza ma lascia acronimi maiuscoli
            title = re.sub(r"([A-Za-z])([A-Za-z']*)", lambda m: m.group(1).upper()+m.group(2).lower(), title)
            title = re.sub(r"\b(AI|LLM|API|SaaS|SQL|CPU|GPU|UI|UX|CLI|SDK|API)\b", lambda m: m.group(0).upper(), title)
            t = f"CHAPTER {ch_idx} — {title}"
        cleaned_h.append((level, t))
    return cleaned_h, blocks


# =====================================================================
# COSTRUZIONE DOCX (stile editoriale richiesto)
# =====================================================================
def _set_page(doc: Document, page: str):
    if page not in ("6x9", "8.5x11"):
        raise ValueError("Formato pagina non valido. Usa '6x9' o '8.5x11'.")
    width_in, height_in = (6.0, 9.0) if page == "6x9" else (8.5, 11.0)
    s = doc.sections[0]
    s.page_width = Inches(width_in)
    s.page_height = Inches(height_in)
    margin = Cm(2.5)
    s.top_margin = margin
    s.bottom_margin = margin
    s.left_margin = margin
    s.right_margin = margin

def _footer_page_numbers(doc: Document):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r._r.append(fld)

def _p(doc: Document, text: str, size: int, bold=False, italic=False, center=False):
    par = doc.add_paragraph()
    if center:
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = par.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bool(bold)
    run.italic = bool(italic)
    return par

def _body_par(doc: Document, text: str):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = par.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return par

def _add_title_page(doc: Document, title: str, subtitle: str, author: str):
    # Spazio ampio sopra
    for _ in range(12):
        doc.add_paragraph("")

    _p(doc, title, size=24, bold=True, center=True)
    if subtitle.strip():
        _p(doc, subtitle, size=12, center=True)

    for _ in range(16):
        doc.add_paragraph("")

    # Logo placeholder: manteniamo solo autore/brand in basso centrato
    _p(doc, author, size=11, italic=False, center=True)
    doc.add_page_break()

def _add_copyright_page(doc: Document, author: str, publisher: Optional[str] = None):
    year = datetime.datetime.now().year
    _p(doc, "Copyright", size=11, bold=True, center=False)
    text = f"© {year} {publisher or author} - All Rights Reserved."
    _body_par(doc, text)
    _body_par(doc, "No portion of this publication may be reproduced, redistributed, or commercially embedded in derivative products without prior authorization.")
    _body_par(doc, "Short excerpts may be used for review or commentary with proper attribution.")
    doc.add_page_break()

def _add_toc_field(doc: Document):
    # Titolo centrato
    _p(doc, "TABLE OF CONTENTS", size=13, bold=True, center=True)
    # Campo TOC aggiornabile
    p = doc.add_paragraph()
    r = p.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-2" \\h \\z \\u')
    r._r.append(fld)
    doc.add_page_break()

def _add_heading(doc: Document, txt: str, level: int):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
    try:
        par.style = doc.styles[f"Heading {level}"]
        run = par.add_run(txt)
        run.font.name = "Arial"
        run.font.size = Pt(12 if level == 1 else 11)
        run.bold = True
    except Exception:
        run = par.add_run(txt)
        run.font.name = "Arial"
        run.font.size = Pt(12 if level == 1 else 11)
        run.bold = True

def build_docx(title: str, subtitle: str, author: str, page: str,
               headings: List[HeadingItem], blocks: List[ContentBlock]) -> Document:
    doc = Document()
    _set_page(doc, page)
    _footer_page_numbers(doc)

    _add_title_page(doc, title, subtitle, author)
    _add_copyright_page(doc, author)
    _add_toc_field(doc)

    bi = 0
    for level, text in headings:
        _add_heading(doc, text, level)
        if bi < len(blocks):
            b_level, paragraphs = blocks[bi]
            if b_level == level:
                for t in paragraphs:
                    _body_par(doc, t)
                bi += 1
        if level == 1:
            doc.add_page_break()
    return doc


# =====================================================================
# COSTRUZIONE PDF (ReportLab) con stile richiesto
# =====================================================================
def _pagesize(page: str):
    return (6.0*inch, 9.0*inch) if page == "6x9" else (8.5*inch, 11.0*inch)

class TOCDoc(SimpleDocTemplate):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)

    def afterFlowable(self, flowable):
        # intercetta titoli per popolare la TOC
        if isinstance(flowable, Paragraph):
            style = flowable.style.name
            txt = flowable.getPlainText()
            if style == "H1":
                self.notify("TOCEntry", (0, txt, self.page))
            elif style == "H2":
                self.notify("TOCEntry", (1, txt, self.page))

def build_pdf(out_pdf: str, title: str, subtitle: str, author: str, page: str,
              headings: List[HeadingItem], blocks: List[ContentBlock]):

    pagesize = _pagesize(page)
    lm = rm = tm = bm = 2.5*cm

    doc = TOCDoc(out_pdf, pagesize=pagesize,
                 leftMargin=lm, rightMargin=rm, topMargin=tm, bottomMargin=bm)

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleCenter", parent=styles["Title"], alignment=1, fontName="Helvetica", fontSize=24, leading=28))
    styles.add(ParagraphStyle(name="Subtitle", parent=styles["Normal"], alignment=1, fontName="Helvetica", fontSize=12, leading=15))
    styles.add(ParagraphStyle(name="Brand", parent=styles["Normal"], alignment=1, fontName="Helvetica", fontSize=11))
    styles.add(ParagraphStyle(name="CopyrightTitle", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=11, spaceAfter=6))
    styles.add(ParagraphStyle(name="CopyrightBody", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=14, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1", parent=styles["Heading1"], fontName="Helvetica-Bold", fontSize=12, spaceBefore=12, spaceAfter=6, leading=15))
    styles.add(ParagraphStyle(name="H2", parent=styles["Heading2"], fontName="Helvetica-Bold", fontSize=11, spaceBefore=8, spaceAfter=4, leading=14))
    styles.add(ParagraphStyle(name="Body", parent=styles["Normal"], fontName="Helvetica", fontSize=11, leading=15, spaceAfter=6))
    styles.add(ParagraphStyle(name="TOCTitle", parent=styles["Normal"], fontName="Helvetica-Bold", alignment=1, fontSize=13, spaceAfter=10))

    story: List = []

    # Pagina titolo (molto bianco sopra)
    story += [Spacer(1, pagesize[1]*0.38),
              Paragraph(title, styles["TitleCenter"])]
    if subtitle.strip():
        story += [Spacer(1, 6), Paragraph(subtitle, styles["Subtitle"])]
    story += [Spacer(1, pagesize[1]*0.28),
              Paragraph(author, styles["Brand"]),
              PageBreak()]

    # Copyright
    year = datetime.datetime.now().year
    story += [
        Paragraph("Copyright", styles["CopyrightTitle"]),
        Paragraph(f"© {year} {author} - All Rights Reserved.", styles["CopyrightBody"]),
        Paragraph("No portion of this publication may be reproduced or redistributed without prior authorization.", styles["CopyrightBody"]),
        Paragraph("Short excerpts may be used with proper attribution.", styles["CopyrightBody"]),
        PageBreak()
    ]

    # TOC
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle(name="TOCLevel1", fontName="Helvetica-Bold", fontSize=11, leftIndent=0, firstLineIndent=0, spaceBefore=8, leading=14, endDots=" ."),
        ParagraphStyle(name="TOCLevel2", fontName="Helvetica", fontSize=10, leftIndent=16, firstLineIndent=0, spaceBefore=2, leading=12, endDots=" ."),
    ]
    story += [Paragraph("TABLE OF CONTENTS", styles["TOCTitle"]), toc, PageBreak()]

    # Contenuti
    bi = 0
    for level, text in headings:
        style = styles["H1"] if level == 1 else styles["H2"]
        story.append(Paragraph(text, style))
        if bi < len(blocks):
            b_level, paragraphs = blocks[bi]
            if b_level == level:
                for t in paragraphs:
                    story.append(Paragraph(t, styles["Body"]))
                bi += 1
        if level == 1:
            story.append(PageBreak())

    # Numeri di pagina centrati in basso
    def on_page(canvas, _doc):
        canvas.saveState()
        w, _ = pagesize
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(w/2.0, 1.0*cm, str(_doc.page))
        canvas.restoreState()

    doc.build(story, onFirstPage=on_page, onLaterPages=on_page)


# =====================================================================
# UI STREAMLIT
# =====================================================================
def ui_main():
    import tempfile
    import streamlit as st

    st.set_page_config(page_title="Formattatore Libro", page_icon="📘", layout="centered")
    st.title("Formattatore Libro")
    st.caption("DOCX + PDF con stile editoriale (senza LibreOffice).")

    with st.sidebar:
        st.header("Impostazioni")
        page = st.selectbox("Formato pagina", ("8.5x11", "6x9"), index=0)
        title = st.text_input("Titolo", "")
        subtitle = st.text_input("Sottotitolo", "")
        author = st.text_input("Autore / Brand", "")
        cleaning = st.selectbox("Pulizia struttura", ("none", "heuristic"), index=1)
        gen_pdf = st.checkbox("Genera anche PDF", value=True)

    up = st.file_uploader("Carica sorgente .docx o .pdf", type=["docx", "pdf"])
    if st.button("Genera"):
        if not up:
            st.error("Carica un file.")
            return
        if not title or not author:
            st.error("Inserisci Titolo e Autore.")
            return

        with tempfile.TemporaryDirectory() as tmp:
            suffix = ".docx" if up.name.lower().endswith(".docx") else ".pdf"
            in_path = os.path.join(tmp, "input" + suffix)
            with open(in_path, "wb") as f:
                f.write(up.getbuffer())

            try:
                if suffix == ".docx":
                    heads, blocks = parse_docx_input(in_path)
                else:
                    heads, blocks = parse_pdf_input(in_path)

                if cleaning == "heuristic":
                    heads, blocks = clean_structure(heads, blocks)

                # DOCX
                st.info("Costruzione DOCX…")
                doc = build_docx(title, subtitle, author, page, heads, blocks)
                out_docx = os.path.join(tmp, "output.docx")
                doc.save(out_docx)
                with open(out_docx, "rb") as f:
                    st.download_button("Scarica DOCX", f.read(),
                        file_name=f"{title.strip().replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

                # PDF
                if gen_pdf:
                    st.info("Costruzione PDF…")
                    out_pdf = os.path.join(tmp, "output.pdf")
                    build_pdf(out_pdf, title, subtitle, author, page, heads, blocks)
                    with open(out_pdf, "rb") as f:
                        st.download_button("Scarica PDF", f.read(),
                            file_name=f"{title.strip().replace(' ','_')}.pdf",
                            mime="application/pdf")

                st.success("Fatto.")

            except Exception as e:
                st.error(f"Errore: {e}")


# =====================================================================
# CLI
# =====================================================================
def cli_main():
    ap = argparse.ArgumentParser(description="Genera DOCX + PDF stile editoriale da DOCX/PDF.")
    ap.add_argument("--input", required=True)
    ap.add_argument("--title", required=True)
    ap.add_argument("--subtitle", default="")
    ap.add_argument("--author", required=True)
    ap.add_argument("--page", choices=["6x9","8.5x11"], default="8.5x11")
    ap.add_argument("--out-docx", default="output.docx")
    ap.add_argument("--out-pdf", default="output.pdf")
    ap.add_argument("--no-pdf", action="store_true")
    ap.add_argument("--clean", choices=["none","heuristic"], default="heuristic")
    args = ap.parse_args()

    src = args.input.lower()
    if src.endswith(".docx"):
        heads, blocks = parse_docx_input(args.input)
    elif src.endswith(".pdf"):
        heads, blocks = parse_pdf_input(args.input)
    else:
        raise ValueError("Formato input non supportato (usa .docx o .pdf)")

    if args.clean == "heuristic":
        heads, blocks = clean_structure(heads, blocks)

    doc = build_docx(args.title, args.subtitle, args.author, args.page, heads, blocks)
    doc.save(args.out_docx)
    print(f"[OK] DOCX: {args.out_docx}")

    if not args.no_pdf:
        build_pdf(args.out_pdf, args.title, args.subtitle, args.author, args.page, heads, blocks)
        print(f"[OK] PDF: {args.out_pdf}")


# =====================================================================
# ENTRY
# =====================================================================
if __name__ == "__main__":
    # Se lanciato con "streamlit run", non passano argomenti -> parte UI
    wants_cli = any(a.startswith("--") for a in sys.argv[1:])
    if wants_cli:
        cli_main()
    else:
        try:
            import streamlit  # noqa
            ui_main()
        except Exception:
            print("UI: streamlit run book_formatter.py")
            print("CLI: python book_formatter.py --input in.docx --title 'T' --author 'A' --page 8.5x11 --clean heuristic")
            sys.exit(1)
